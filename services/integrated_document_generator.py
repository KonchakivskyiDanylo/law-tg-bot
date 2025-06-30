# integrated_document_generator.py
import re
import json
import os
import logging
from typing import Dict, List, Optional
from datetime import datetime
from openai import AsyncOpenAI
from telegram import Bot
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import docx.oxml
from abc import ABC, abstractmethod
from docx.enum.table import WD_TABLE_ALIGNMENT
from services.subscription_service import get_conversation_history, append_to_last_request_dialog, \
    start_new_request_session

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

from prompts import (
    CONTRACT_PROMPT, APPLICATION_PROMPT, ACT_PROMPT, CLAIM_PROMPT,
    POWER_OF_ATTORNEY_PROMPT, PRETENSE_PROMPT, DOCUMENT_PROMPTS, DOCUMENT_COMPLETENESS_EVALUATION_PROMPT
)

from config.config import OPENAI_API_KEY, MAX_TELEGRAM_MESSAGE_LENGTH

client = AsyncOpenAI(api_key=OPENAI_API_KEY)


class IntegratedDocumentGenerator:
    def __init__(self):
        self.client = client
        self.conversation_history = []

    async def get_specialized_prompt(self, document_type: str) -> str:
        return DOCUMENT_PROMPTS.get(document_type, CONTRACT_PROMPT)

    async def evaluate_document_completeness(self, messages: List[Dict[str, str]], document_type: str) -> dict:
        """
        Evaluate the completeness of a document based on provided messages and document type.

        This asynchronous method analyzes the provided messages to evaluate the completeness
        of a document request. It delegates the evaluation process to an OpenAI completion
        function, which processes the input data and generates a response in the form of a
        structured object. The method also handles errors in processing or decoding the response
        and ensures a fallback result is returned in such cases.

        :param messages: A list of dictionaries representing the conversation messages. Each
            dictionary contains `role` and `content` keys.
        :param document_type: A string indicating the type of document being evaluated.
        :return: A dictionary that includes the completeness score, clarifying questions,
            an explanation, the provided document type, and any missing information.
        """
        logger.info(f"Evaluating document completeness for type: {document_type}")

        try:
            evaluation_messages = messages.copy()
            evaluation_messages.append({
                "role": "system",
                "content": f"Type of document to create: {document_type}"
            })

            response = await self.create_openai_completion(
                messages=[
                    {"role": "system", "content": DOCUMENT_COMPLETENESS_EVALUATION_PROMPT},
                    {"role": "user", "content": f"Запрос пользователя: {messages[-1].get('content', '')}"}
                ],
                response_format={"type": "json_object"}
            )

            if not response:
                logger.error("No response from completeness evaluation")
                return {
                    "score": 1,
                    "clarifying_questions": ["Please provide more information to create the document."],
                    "explanation": "Error analyzing the request.",
                    "document_type": document_type,
                    "missing_info": ["Request details"]
                }

            try:
                result = json.loads(response)
                logger.info(f"Completeness evaluation result: score={result.get('score')}")
                return result
            except json.JSONDecodeError as e:
                logger.error(f"JSON decode error in completeness evaluation: {e}")
                return {
                    "score": 1,
                    "clarifying_questions": ["Please provide more information to create the document."],
                    "explanation": "Error analyzing the request.",
                    "document_type": document_type,
                    "missing_info": ["Request details"]
                }

        except Exception as e:
            logger.error(f"Error in evaluate_document_completeness: {str(e)}")
            return {
                "score": 1,
                "clarifying_questions": ["Please provide more information to create the document."],
                "explanation": "Error analyzing the request.",
                "document_type": document_type,
                "missing_info": ["Request details"]
            }

    async def generate_document_with_completeness_check(self, user_request: str, document_type: str = "",
                                                        conversation_messages: List[Dict[str, str]] = None) -> dict:
        """
        This asynchronous method generates a document based on the user request and performs a completeness
        check to ensure all necessary information is provided. If the evaluation determines that the
        request lacks sufficient details, clarifying questions are returned to the user for additional input.
        The method utilizes GPT-related functions to analyze and evaluate completeness, as well as to
        generate the document when adequate information is available.

        :param user_request: The input string containing the user's request for document generation.
        :param document_type: The type of document to generate. Defaults to an empty string.
                              If not specified, it will be inferred using a GPT function.
        :param conversation_messages: A list of conversation messages in the form of dictionaries, with
                                      "role" and "content" as keys. Used during the document generation
                                      and evaluation process. Defaults to None.

        :return: A dictionary containing the status of the operation and relevant messages. If completeness
                 requirements are unmet, additional clarifying questions, missing information details, and
                 an explanation are provided. If successful, the generated document is returned.
        """
        try:
            if conversation_messages is None:
                conversation_messages = [{"role": "user", "content": user_request}]

            if not document_type:
                document_type = await get_document_type_gpt(user_request)

            logger.info(f"Generating document of type: {document_type}")

            evaluation = await self.evaluate_document_completeness(conversation_messages, document_type)

            if evaluation.get("score", 0) < 4:
                clarifying_questions = evaluation.get("clarifying_questions", [
                    "Please, give more information to create document."
                ])
                explanation = evaluation.get("explanation",
                                             "Not enough information.")
                missing_info = evaluation.get("missing_info", [])

                response_text = f"To create a high-quality document of type «{document_type}», I need additional information:\n\n"

                for i, question in enumerate(clarifying_questions, 1):
                    response_text += f"{i}. {question}\n"

                if missing_info:
                    response_text += f"\nIt is especially important to clarify:\n"
                    for info in missing_info:
                        response_text += f"• {info}\n"

                response_text += f"\n{explanation}"

                return {
                    "status": "incomplete",
                    "message": response_text,
                    "document_type": document_type,
                    "needs_clarification": True,
                    "clarifying_questions": clarifying_questions,
                    "missing_info": missing_info
                }

            return await self.generate_document_with_specialized_prompt(user_request, document_type,
                                                                        conversation_messages)

        except Exception as e:
            logger.error(f"Error in generate_document_with_completeness_check: {e}")
            return {
                "status": "error",
                "message": f"Error: {str(e)}",
                "document_type": document_type
            }

    async def generate_document_with_specialized_prompt(self, user_request: str, document_type: str = "",
                                                        conversation_messages: List[Dict[str, str]] = None) -> dict:
        """
        Handles the generation of a document using a specialized prompt tailored to the request
        and context provided. This method utilizes a combination of a predefined specialized
        prompt and the conversation's context to create a JSON-formatted document based on the
        user request.

        :param user_request: The specific request made by the user that serves as the basis
            for the document generation.
        :type user_request: str
        :param document_type: The type of document being requested. Defaults to an empty
            string if not provided.
        :type document_type: str
        :param conversation_messages: A list of messages representing the context of
            the conversation. Each message is a dictionary with "role" and "content" keys.
            This is optional and may be None.
        :type conversation_messages: List[Dict[str, str]]
        :return: A dictionary containing the status of the operation, the document type,
            generated document text in JSON, and whether clarification is needed.
        :rtype: dict
        """
        try:
            specialized_prompt = await self.get_specialized_prompt(document_type)

            context = ""
            if conversation_messages and len(conversation_messages) > 1:
                context = "Conversation context:\n"
                for msg in conversation_messages[:-1]:
                    role = "User" if msg["role"] == "user" else "Assistant"
                    context += f"{role}: {msg['content']}\n"
                context += "\n"

            full_prompt = f"""
            {specialized_prompt}

            {context} User request: {user_request}

            Create the document strictly in JSON format without any additional comments.
            All fields must be filled with relevant information based on the request and context.
            Use all available information from the conversation context.
            """

            response = await self.create_openai_completion(
                messages=[{"role": "user", "content": full_prompt}],
                response_format={"type": "json_object"}
            )

            if not response:
                return {"status": "error", "message": "Failed to get a response from AI"}

            try:
                document_data = json.loads(response)
                return {
                    "status": "success",
                    "document_type": document_data.get("document_type", document_type),
                    "document_text": document_data,
                    "needs_clarification": False
                }
            except json.JSONDecodeError as e:
                logger.error(f"Error with JSON: {e}")
                return {"status": "error", "message": "Error with AI"}

        except Exception as e:
            logger.error(f"Error in generate_document_with_specialized_prompt: {e}")
            return {"status": "error", "message": str(e)}

    async def create_openai_completion(self, messages: List[Dict[str, str]], response_format: Dict[str, str] = None) -> \
            Optional[str]:
        """
        Sends a request to the OpenAI API to create a chat completion using the provided
        messages and optional response format. This function leverages the GPT-4.1 model
        and allows for a maximum of 16,000 tokens in the completion. It processes the
        response from the API, and if valid, extracts and returns the content of the
        first message choice. If any error occurs or the response is invalid, it
        returns None.

        :param messages: A list of dictionaries, where each dictionary represents a
            message to be sent to the OpenAI API. Each message must include a 'role'
            (e.g., 'user', 'assistant') and 'content'.
        :type messages: List[Dict[str, str]]
        :param response_format: An optional dictionary specifying the response format
            to be applied to the OpenAI API output.
        :type response_format: Dict[str, str], optional
        :return: The content of the first message choice from the API response if the
            request is successful and valid, otherwise None.
        :rtype: Optional[str]
        """
        try:
            logger.info(f"Sending request to OpenAI API with {len(messages)} messages")

            params = {
                "model": "gpt-4.1",
                "messages": messages,
                "max_completion_tokens": 16000,
            }

            if response_format:
                params["response_format"] = response_format

            response = await self.client.chat.completions.create(**params)

            if response and response.choices and len(response.choices) > 0:
                return response.choices[0].message.content
            else:
                logger.error("Invalid response from OpenAI API")
                return None

        except Exception as e:
            logger.error(f"Error calling OpenAI API: {str(e)}", exc_info=True)
            return None


class DocumentTemplate(ABC):
    def __init__(self, doc: Document):
        self.doc = doc
        self.setup_styles()
        self.setup_page_format()

    def setup_styles(self):
        styles = self.doc.styles

        if "Normal" in styles:
            normal_style = styles["Normal"]
            normal_style.font.size = Pt(12)
            normal_style.font.name = "Times New Roman"
            normal_style.paragraph_format.space_after = Pt(6)
            normal_style.paragraph_format.line_spacing = 1.5
            normal_style.paragraph_format.first_line_indent = Cm(1.25)

        if "Title" in styles:
            title_style = styles["Title"]
            title_style.font.size = Pt(14)
            title_style.font.bold = True
            title_style.font.name = "Times New Roman"

    def setup_page_format(self):
        section = self.doc.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    def add_paragraph(self, text, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, bold=False,
                      underline=False, indent_first_line=True, space_after=6):
        p = self.doc.add_paragraph()
        run = p.add_run(text)

        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.bold = bold
        run.underline = underline

        p.alignment = alignment
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(space_after)

        if indent_first_line and alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
            p.paragraph_format.first_line_indent = Cm(1.25)

        return p

    def add_centered_number(self, number_text):
        p = self.doc.add_paragraph()
        run = p.add_run(number_text)

        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.bold = True

        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.space_before = Pt(12)

        return p

    def add_table_row(self, left_text="", right_text="", left_bold=False, right_bold=False):
        table = self.doc.add_table(rows=1, cols=2)

        for row in table.rows:
            for cell in row.cells:
                cell._element.get_or_add_tcPr().append(docx.oxml.parse_xml(
                    r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    r'<w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/>'
                    r'</w:tcBorders>'))

        if left_text:
            left_cell = table.cell(0, 0)
            left_para = left_cell.paragraphs[0]
            left_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            left_run = left_para.add_run(left_text)
            left_run.font.name = "Times New Roman"
            left_run.font.size = Pt(14)
            left_run.bold = left_bold

        if right_text:
            right_cell = table.cell(0, 1)
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            right_run = right_para.add_run(right_text)
            right_run.font.name = "Times New Roman"
            right_run.font.size = Pt(14)
            right_run.bold = right_bold

    def add_parties_info(self, sender_info, recipient_info):
        if sender_info or recipient_info:
            table = self.doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            for row in table.rows:
                for cell in row.cells:
                    cell._element.get_or_add_tcPr().append(docx.oxml.parse_xml(
                        r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                        r'<w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/>'
                        r'</w:tcBorders>'))

            if sender_info:
                left_cell = table.cell(0, 0)
                left_para = left_cell.paragraphs[0]
                left_run = left_para.add_run(sender_info)
                left_run.font.name = "Times New Roman"
                left_run.font.size = Pt(14)
                left_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            if recipient_info:
                right_cell = table.cell(0, 1)
                right_para = right_cell.paragraphs[0]
                right_run = right_para.add_run(recipient_info)
                right_run.font.name = "Times New Roman"
                right_run.font.size = Pt(14)
                right_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    @abstractmethod
    async def generate(self, json_data: dict):
        pass


class ContractTemplate(DocumentTemplate):
    def _is_main_section_number(self, number_str):
        return bool(re.match(r'^\d+$', str(number_str).strip()))

    def _process_numbered_content(self, content_list):
        if not content_list:
            return

        for item in content_list:
            if isinstance(item, dict):
                number = item.get("number", "")
                text = item.get("text", "")
                level = item.get("level", 0)
                subitems = item.get("subitems", [])

                if self._is_main_section_number(number):
                    if text:
                        full_title = f"{number}.{text.upper()}"
                    else:
                        full_title = f"{number}."

                    self.add_centered_number(full_title)
                else:
                    full_text = f"{number}. {text}" if number and text else (text or f"{number}.")

                    if full_text:
                        self.add_paragraph(
                            full_text,
                            alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                            indent_first_line=True
                        )
                if subitems:
                    self._process_numbered_content(subitems)
            else:
                self.add_paragraph(
                    str(item),
                    alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                    indent_first_line=True
                )

    async def generate(self, json_data: dict):
        if json_data.get("document_title"):
            title_text = json_data["document_title"].upper()
            self.add_paragraph(
                title_text,
                alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
                bold=True,
                indent_first_line=False,
                space_after=12
            )
        if json_data.get("city") or json_data.get("date_place"):
            self.add_table_row(
                left_text=f"{json_data.get('city', '')}" if json_data.get("city") else "",
                right_text=json_data.get("date_place", "")
            )
        if json_data.get("heading"):
            self.add_paragraph(
                json_data["heading"],
                alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                indent_first_line=True
            )
        if json_data.get("introduction"):
            self.add_paragraph(
                json_data["introduction"],
                alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                indent_first_line=True
            )

        if json_data.get("main_body"):
            self._process_numbered_content(json_data["main_body"])

        if json_data.get("conclusion"):
            self.add_paragraph(
                json_data["conclusion"],
                alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                indent_first_line=True
            )

        if json_data.get("appendices"):
            self.doc.add_paragraph()
            self.add_paragraph("Application:", bold=True, indent_first_line=False)
            for item in json_data["appendices"]:
                self.add_paragraph(f"— {item}", alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, indent_first_line=False)

        if json_data.get("executor_info"):
            self.doc.add_paragraph()
            self.add_paragraph(
                f"Executor: {json_data['executor_info']}",
                alignment=WD_PARAGRAPH_ALIGNMENT.LEFT,
                indent_first_line=False
            )

        if json_data.get("parties_details"):
            self.add_paragraph("Party details", alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
                               bold=True)
            self.add_parties_info(json_data.get("parties_details"), json_data.get("parties_details2"))
            self.doc.add_paragraph()

        if json_data.get("signatures"):
            self.doc.add_paragraph()
            self.add_paragraph("Signatures", alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
                               bold=True)

            signatures = json_data["signatures"]
            sender_sig = signatures.get("sender", {})
            recipient_sig = signatures.get("recipient", {})

            sender_lines = []
            recipient_lines = []

            if sender_sig:
                label = sender_sig.get("label", "")

                if label:
                    sender_lines.append(label)

            if recipient_sig:
                label = recipient_sig.get("label", "")
                if label:
                    recipient_lines.append(label)
            if sender_lines and recipient_lines:
                max_lines = max(len(sender_lines), len(recipient_lines))
                for i in range(max_lines):
                    left_text = sender_lines[i] if i < len(sender_lines) else ""
                    right_text = recipient_lines[i] if i < len(recipient_lines) else ""
                    self.add_table_row(left_text=left_text, right_text=right_text)
            elif sender_lines:
                for line in sender_lines:
                    self.add_paragraph(line, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, indent_first_line=False)
            elif recipient_lines:
                for line in recipient_lines:
                    self.add_paragraph(line, alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT, indent_first_line=False)
        if json_data.get("stamp_area"):
            self.doc.add_paragraph()
            self.add_paragraph("(SEAL)", alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT, indent_first_line=False)


class ApplicationTemplate(DocumentTemplate):
    async def generate(self, json_data: dict):
        if json_data.get("recipient"):
            self.add_paragraph(json_data["recipient"], alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT)

        if json_data.get("sender"):
            self.add_paragraph(f"от {json_data['sender']}", alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT)

        self.doc.add_paragraph()

        title = json_data.get("document_title")
        self.add_paragraph(title.upper(), alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)

        self.doc.add_paragraph()

        if json_data.get("main_body"):
            for paragraph in json_data["main_body"]:
                self.add_paragraph(paragraph, alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

        self.doc.add_paragraph()
        if json_data.get("date_place"):
            self.add_table_row(
                left_text=json_data["date_place"],
                right_text="________________"
            )


class ComplaintTemplate(DocumentTemplate):
    async def generate(self, json_data: dict):
        if json_data.get("recipient"):
            self.add_paragraph(json_data["recipient"], alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT)

        if json_data.get("sender"):
            self.add_paragraph(f"{json_data['sender']}", alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT)

        if json_data.get("document_title"):
            self.add_paragraph(json_data["document_title"].upper(), alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)

        if json_data.get("main_body"):
            self.add_paragraph(json_data["main_body"], alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

        if json_data.get("appendices"):
            self.add_paragraph("Appendices:", bold=True)
            for item in json_data["appendices"]:
                self.add_paragraph(f"- {item}", alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)

        if json_data.get("date_place"):
            self.add_paragraph(json_data["date_place"], alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)


class OrderTemplate(DocumentTemplate):
    async def generate(self, json_data: dict):
        if json_data.get("organisation_name"):
            self.add_paragraph(json_data["organisation_name"], alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)
        if json_data.get("document_name"):
            self.add_paragraph(json_data["document_name"].upper(), alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
        if json_data.get("document_number") or json_data.get("date_place"):
            self.add_table_row(
                left_text=f"{json_data.get('city', '')}" if json_data.get("document_number") else "",
                right_text=json_data.get("date_place", "")
            )
        if json_data.get("document_title"):
            self.add_paragraph(json_data["document_title"].upper(), alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
        if json_data.get("main_body"):
            self.add_paragraph(json_data["main_body"], alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

        if json_data.get("appendices"):
            self.add_paragraph("Appendices:", bold=True)
            for item in json_data["appendices"]:
                self.add_paragraph(f"- {item}", alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)

        if json_data.get("sender_workplace") or json_data.get("sender_name"):
            self.add_table_row(
                left_text=f"{json_data.get('sender_workplace', '')}" if json_data.get("sender_workplace") else "",
                right_text=json_data.get("sender_name", "")
            )
        if json_data.get("recipients"):
            self.add_paragraph(json_data["recipients"], alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)


class ClaimTemplate(DocumentTemplate):
    async def generate(self, json_data: dict):

        if json_data.get("court"):
            self.add_paragraph(json_data["court"], alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT)

        if json_data.get("plaintiff"):
            self.add_paragraph("Plaintiff:", bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT)
            self.add_paragraph(json_data["plaintiff"], alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT)

        if json_data.get("defendant"):
            self.add_paragraph("Defendant:", bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT)
            self.add_paragraph(json_data["defendant"], alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT)
            self.add_paragraph("")

        if json_data.get("document_title"):
            self.add_paragraph(json_data["document_title"].upper(), alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)

        if json_data.get("circumstances"):
            self.add_paragraph(json_data["circumstances"], alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

        if json_data.get("legal_basis"):
            self.add_paragraph(json_data["legal_basis"], alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

        if json_data.get("petition"):
            self.add_paragraph("Petition:", bold=True)
            for paragraph in json_data["petition"]:
                self.add_paragraph(paragraph, alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
            self.add_paragraph("")

        if json_data.get("date_place"):
            self.add_paragraph(json_data["date_place"], alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
            self.add_paragraph("")

        if json_data.get("signature"):
            self.add_paragraph(json_data["signature"], alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)


class PretenseTemplate(DocumentTemplate):
    def _process_numbered_content(self, content_list):
        if not content_list:
            return

        for item in content_list:
            if isinstance(item, dict):
                text = item.get("text", "")
                number = item.get("number", "")
                level = item.get("level", 0)
                subitems = item.get("subitems", [])

                indent_levels = [0, 0.5, 1.0, 1.5, 2.0]

                if number:
                    full_text = f"{number}. {text}" if text else f"{number}."
                else:
                    full_text = text

                if full_text:
                    self.add_paragraph(
                        full_text,
                        alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                    )

                if subitems:
                    self._process_numbered_content(subitems)

            else:
                self.add_paragraph(str(item), alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

    async def generate(self, json_data: dict):
        if json_data.get("sender_info") or json_data.get("recipient_info"):
            sender_info = json_data.get("sender_info", {})
            recipient_info = json_data.get("recipient_info", {})

            sender_text = ""
            if sender_info:
                sender_parts = []
                if sender_info.get("organization"):
                    sender_parts.append(sender_info["organization"])
                if sender_info.get("address"):
                    sender_parts.append(f"Address: {sender_info['address']}")
                if sender_info.get("postal_address"):
                    sender_parts.append(f"Address: {sender_info['postal_address']}")
                if sender_info.get("phone"):
                    sender_parts.append(f"Phone: {sender_info['phone']}")
                if sender_info.get("email"):
                    sender_parts.append(f"www.{sender_info['email']}")

                sender_text = "\n".join(sender_parts)

            recipient_text = ""
            if recipient_info:
                recipient_parts = []
                if recipient_info.get("position"):
                    recipient_parts.append(recipient_info["position"])
                if recipient_info.get("name"):
                    recipient_parts.append(recipient_info["name"])
                if recipient_info.get("address"):
                    recipient_parts.append(recipient_info["address"])
                if recipient_info.get("copy_to"):
                    recipient_parts.append("copy")
                    recipient_parts.append(recipient_info["copy_to"])

                recipient_text = "\n".join(recipient_parts)

            self.add_table_row(
                left_text=sender_text,
                right_text=recipient_text
            )

        if json_data.get("document_title"):
            self.add_paragraph("")
            self.add_paragraph(
                json_data["document_title"].upper(),
                alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
                bold=True
            )
        if json_data.get("subtitle"):
            self.add_paragraph(
                json_data["subtitle"],
                alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
            )
        if json_data.get("date") or json_data.get("document_number"):
            date_text = json_data.get("date", "")
            number_text = ""
            if json_data.get("document_number"):
                number_text = f"№{json_data['document_number']}"

            self.add_paragraph("")
            self.add_paragraph(
                f"{date_text} {number_text}".strip(),
                alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            )
        if json_data.get("claim_text"):
            self.add_paragraph("")
            self.add_paragraph(
                json_data["claim_text"],
                alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            )
        if json_data.get("obligations"):
            self._process_numbered_content(json_data["obligations"])

        if json_data.get("additional_text"):
            for paragraph in json_data["additional_text"]:
                self.add_paragraph(paragraph, alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
        if json_data.get("attachments"):
            self.add_paragraph("")
            self.add_paragraph("Attachments:", bold=True)
            for attachment in json_data["attachments"]:
                self.add_paragraph(f"- {attachment}")

        if json_data.get("signatures"):
            self.add_paragraph("")
            self.add_paragraph("signatures")


class LetterTemplate(DocumentTemplate):
    async def generate(self, json_data: dict):
        if json_data.get("recipient"):
            self.add_paragraph(json_data["recipient"], alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT)
            self.add_paragraph("")
            self.add_paragraph("")

        if json_data.get("Intro"):
            self.add_paragraph(json_data["Intro"], alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
            self.add_paragraph("")

        if json_data.get("main_body"):
            self.add_paragraph(json_data["main_body"], alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
            self.add_paragraph("")
            self.add_paragraph("")

        if json_data.get("sender"):
            self.add_paragraph(json_data["sender"], alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)


class PowerOfAttorneyTemplate(DocumentTemplate):
    async def generate(self, json_data: dict):
        if json_data.get("document_type"):
            self.add_paragraph(json_data["document_type"].upper(), alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
            self.add_paragraph(json_data["document_desc"], alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
            self.add_paragraph("")

        if json_data.get("date") or json_data.get("place"):
            self.add_table_row(
                left_text=f"{json_data.get('place', '')}" if json_data.get("place") else "",
                right_text=json_data.get("date", "")
            )
            self.doc.add_paragraph()
            self.add_paragraph("")

        if json_data.get("main_body"):
            self.add_paragraph(json_data["main_body"], alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
        if json_data.get("validity_period"):
            self.add_paragraph(json_data["validity_period"], alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
            self.doc.add_paragraph()
            self.doc.add_paragraph()

        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        table.allow_autofit = True

        cell = table.cell(0, 0)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        run = paragraph.add_run("________________________ / ________________________\n«___» _____________ 2025 г.")
        run.font.size = Pt(10)


class ReportTemplate(DocumentTemplate):
    async def generate(self, json_data: dict):
        if json_data.get("organization_name"):
            self.add_paragraph(json_data["organization_name"], alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)
        address_info = []
        if json_data.get("organization_address"):
            address_info.append(json_data["organization_address"])
        if address_info:
            self.add_paragraph("; ".join(address_info), alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)

        self.doc.add_paragraph()

        if json_data.get("document_title"):
            self.add_paragraph(json_data["document_title"].upper(), alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)

        if json_data.get("report_date"):
            self.add_paragraph(json_data["report_date"], alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)

        self.doc.add_paragraph()
        if json_data.get("legal_basis"):
            self.add_paragraph(json_data["legal_basis"], alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

        if json_data.get("summary_table"):
            table = self.doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "№ п/п"
            hdr_cells[1].text = "Information Name"
            hdr_cells[2].text = "Number of Contracts"
            hdr_cells[3].text = "Contract Amount"

            for item in json_data["summary_table"]:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item.get("number", ""))
                row_cells[1].text = item.get("description", "")
                row_cells[2].text = str(item.get("contracts_count", ""))
                row_cells[3].text = item.get("total_amount", "")

        if json_data.get("executor_info"):
            self.doc.add_paragraph()  # отступ
            self.add_paragraph(f"Executor: {json_data['executor_info']}", alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)

        if json_data.get("appendices"):
            self.add_paragraph("Appendices:", bold=True)
            for item in json_data["appendices"]:
                self.add_paragraph(f"- {item}", alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)

        if json_data.get("stamp_area"):
            self.doc.add_paragraph()
            self.add_paragraph("(М.П.)", alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT, indent_first_line=False)


class ProtocolTemplate(DocumentTemplate):
    async def generate(self, json_data: dict):
        if json_data.get("document_title"):
            self.add_paragraph(json_data["document_title"].upper(), alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)

        if json_data.get("place") or json_data.get("date"):
            number = f" {json_data.get('place', '')}" if json_data.get("place") else ""
            date_place = json_data.get("date", "")
            self.add_table_row(left_text=number, right_text=date_place)

        if json_data.get("participants"):
            self.add_paragraph("Participants:", bold=True)
            for p in json_data["participants"]:
                self.add_paragraph(p)

        if json_data.get("agenda"):
            self.add_paragraph("Agenda:", bold=True)
            for i, item in enumerate(json_data["agenda"], 1):
                self.add_paragraph(f"{i}. {item}")

        if json_data.get("main_body"):
            self.add_paragraph("")
            for section in json_data["main_body"]:
                self.add_paragraph(section, alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

        if json_data.get("signatures"):
            self.doc.add_paragraph()
            self.add_paragraph("Signatures:", bold=True)

            sig = json_data["signatures"]
            left = sig.get("sender", {})
            right = sig.get("recipient", {})

            left_lines = []
            right_lines = []

            if left.get("label"):
                left_lines.append(left["label"])
            if left.get("position"):
                left_lines.append(left["position"])
            if left.get("name"):
                left_lines.append(left["name"])

            if right.get("label"):
                right_lines.append(right["label"])
            if right.get("position"):
                right_lines.append(right["position"])
            if right.get("name"):
                right_lines.append(right["name"])

            max_lines = max(len(left_lines), len(right_lines))
            for i in range(max_lines):
                l = left_lines[i] if i < len(left_lines) else ""
                r = right_lines[i] if i < len(right_lines) else ""
                self.add_table_row(left_text=l, right_text=r)

        # Печать
        if json_data.get("stamp_area"):
            self.add_paragraph()
            self.add_paragraph("(М.П.)", alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT)


class DocumentTemplateFactory:
    _templates = {
        "contract": ContractTemplate,
        "complaint": ComplaintTemplate,
        "order": OrderTemplate,
        "application": ApplicationTemplate,
        "pretense": PretenseTemplate,
        "act": ContractTemplate,
        "claim": ClaimTemplate,
        "objection": ClaimTemplate,
        "power_of_attorney": PowerOfAttorneyTemplate,
        "letter": LetterTemplate,
        "report": ReportTemplate,
        "protocol": ProtocolTemplate
    }

    @classmethod
    def get_template(cls, document_type: str, doc: Document) -> DocumentTemplate:
        template_class = cls._templates.get(document_type.lower())
        if not template_class:
            return ContractTemplate(doc)
        return template_class(doc)


async def save_as_docx(json_data: dict, filename: str = "document.docx") -> str:
    try:
        doc = Document()
        document_type = json_data.get("document_type", "act")
        template = DocumentTemplateFactory.get_template(document_type, doc)
        await template.generate(json_data)

        file_path = os.path.join(os.getcwd(), filename)
        doc.save(file_path)
        logger.info(f"Document saved as {file_path}")
        return file_path

    except Exception as e:
        logger.error(f"Error in saving document: {str(e)}", exc_info=True)
        return None


async def get_document_type_gpt(message: str) -> str:
    from prompts import SYSTEM_PROMPT
    response = await client.chat.completions.create(
        model="o3-mini",
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": message}
        ]
    )

    reply = response.choices[0].message.content.strip().lower()
    match = re.search(r"type:\s*(\w+)", reply)

    if match:
        document_type = match.group(1)
        logger.info(f"[GPT] Document type recognized: {document_type} ← from message: \"{message}\"")
        return document_type
    logger.error(f"[GPT] Unexpected reply format: \"{reply}\"")
    return "act"


async def process_user_message_integrated(message: str, chat_id: int, bot: Bot,
                                          conversation_messages: List[Dict[str, str]] = None) -> tuple[str, str | None]:
    user_id = str(chat_id)

    conversation_messages = get_conversation_history(user_id)

    if not conversation_messages:
        start_new_request_session(user_id, question=message, session_type="document")
        conversation_messages = [{"role": "user", "content": message}]
    else:
        append_to_last_request_dialog(user_id, role="user", message=message)
        conversation_messages.append({"role": "user", "content": message})
    generator = IntegratedDocumentGenerator()

    document_type = await get_document_type_gpt(message)

    result = await generator.generate_document_with_completeness_check(
        user_request=message,
        document_type=document_type,
        conversation_messages=conversation_messages
    )
    if result.get("status") == "incomplete":
        assistant_response = result.get("message", "Additional information is required to create the document.")
        append_to_last_request_dialog(user_id, role="bot", message=assistant_response)
        return truncate_if_needed(assistant_response), None

    elif result.get("status") == "success":
        file_path = await save_as_docx(
            result.get("document_text"),
            f"{result.get('document_type')}_{chat_id}_{int(datetime.now().timestamp())}.docx"
        )
        assistant_response = (
            f"✅ The document «{result.get('document_type')}» was successfully created using the specialized template."
            if file_path else "❌ The document was created, but the file could not be saved."
        )

        append_to_last_request_dialog(user_id, role="bot", message=assistant_response)
        return truncate_if_needed(assistant_response), file_path


    else:
        assistant_response = f"❌ Failed to create the document: {result.get('message', 'Unknown error')}"
        append_to_last_request_dialog(user_id, role="bot", message=assistant_response)
        return truncate_if_needed(assistant_response), None


def truncate_if_needed(text: str) -> str:
    if len(text) > MAX_TELEGRAM_MESSAGE_LENGTH:
        return text[:MAX_TELEGRAM_MESSAGE_LENGTH - 100] + "\n\n[Message truncated]"
    return text
