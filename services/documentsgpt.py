import logging
import json
from docx import Document
from PyPDF2 import PdfReader
from typing import Dict, List, Any, Optional
from openai import AsyncOpenAI
import docx
import docx.oxml
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from telegram import Bot
from config.config import OPENAI_API_KEY, MAX_TELEGRAM_MESSAGE_LENGTH
import os
from prompts import DOCUMENT_GENERATOR_PROMPT_2, DOCUMENT_ANALYSIS_PROMPT, DOCUMENT_GENERATOR_PROMPT, \
    DOCUMENT_TYPE_DETECTION_PROMPT, DOCUMENT_GENERATION_FROM_DIALOGUE_PROMPT, RECOMMENDATIONS_PROMPT, VALIDATION_PROMPT

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

if not OPENAI_API_KEY:
    logger.warning("OPENAI_API_KEY isn't found")

client = AsyncOpenAI(api_key=OPENAI_API_KEY)

conversation_history = []


def save_as_docx(json_data: dict, filename: str = "document.docx") -> str:
    try:
        doc = Document()
        styles = doc.styles

        if "Normal" in styles:
            normal_style = styles["Normal"]
            normal_style.font.size = Pt(12)
            normal_style.font.name = "Times New Roman"
            normal_style.paragraph_format.space_after = Pt(6)
            normal_style.paragraph_format.line_spacing = 1.15
            normal_style.paragraph_format.first_line_indent = Cm(1.25)

        if "Title" in styles:
            title_style = styles["Title"]
            title_style.font.size = Pt(16)
            title_style.font.bold = True
            title_style.font.name = "Times New Roman"

        section = doc.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        run = footer_para.runs[0] if footer_para.runs else footer_para.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        def add_paragraph(text, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, bold=False, underline=False,
                          style_name=None, first_line_indent=None, space_before=None, space_after=None):
            if style_name and style_name in styles:
                p = doc.add_paragraph(text, style=style_name)
            else:
                p = doc.add_paragraph(text)

            p.alignment = alignment

            if first_line_indent is not None:
                p.paragraph_format.first_line_indent = first_line_indent
            if space_before is not None:
                p.paragraph_format.space_before = space_before
            if space_after is not None:
                p.paragraph_format.space_after = space_after

            if p.runs:
                run = p.runs[0]
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run.bold = bold
                run.underline = underline

            return p

        def add_heading(text, level=1):
            if level == 1:
                p = add_paragraph(text, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, bold=True,
                                  space_before=Pt(12), space_after=Pt(6), first_line_indent=Cm(0))
            elif level == 2:
                p = add_paragraph(text, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, bold=True,
                                  space_before=Pt(6), space_after=Pt(3), first_line_indent=Cm(0))
            else:
                p = add_paragraph(text, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, bold=False,
                                  space_before=Pt(3), space_after=Pt(3), first_line_indent=Cm(0))
            return p

        def add_table_row(left_text="", right_text="", left_bold=False, right_bold=False):
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            table.columns[0].width = Cm(10)
            table.columns[1].width = Cm(7)

            for row in table.rows:
                for cell in row.cells:
                    cell._element.get_or_add_tcPr().append(docx.oxml.parse_xml(
                        r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/></w:tcBorders>'))

            if left_text:
                left_cell = table.cell(0, 0)
                left_para = left_cell.paragraphs[0]
                left_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                left_run = left_para.add_run(left_text)
                left_run.font.name = "Times New Roman"
                left_run.font.size = Pt(12)
                left_run.bold = left_bold

            if right_text:
                right_cell = table.cell(0, 1)
                right_para = right_cell.paragraphs[0]
                right_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                right_run = right_para.add_run(right_text)
                right_run.font.name = "Times New Roman"
                right_run.font.size = Pt(12)
                right_run.bold = right_bold

        def determine_heading_level(number):
            if not number:
                return 1
            dots = number.count('.')
            if dots == 1:
                return 1
            elif dots == 2:
                return 2
            else:
                return 3

        if json_data.get("sender") or json_data.get("recipient"):
            add_table_row(
                left_text=json_data.get("sender", ""),
                right_text=json_data.get("recipient", "")
            )
            doc.add_paragraph()

        if json_data.get("organization_name"):
            add_paragraph(json_data["organization_name"], alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
                          bold=True, space_after=Pt(12))

        if json_data.get("document_title"):
            add_paragraph(json_data["document_title"].upper(), alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
                          bold=True, space_after=Pt(18))

        if json_data.get("document_number") or json_data.get("date_place"):
            add_table_row(
                left_text=f"№ {json_data.get('document_number', '')}" if json_data.get("document_number") else "",
                right_text=json_data.get("date_place", "")
            )
            doc.add_paragraph()

        if json_data.get("heading"):
            add_paragraph(json_data["heading"], alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                          space_after=Pt(12))

        if json_data.get("introduction"):
            add_paragraph(json_data["introduction"], alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                          space_after=Pt(12))

        if json_data.get("main_body"):
            for item in json_data["main_body"]:
                if isinstance(item, dict):
                    item_type = item.get("type", "text")
                    number = item.get("number", "")
                    title = item.get("title", "")
                    content = item.get("content", "")

                    if item_type in ["section", "subsection", "point"]:
                        if title:
                            header_text = f"{number} {title}"
                        else:
                            header_text = number

                        if header_text.strip():
                            level = determine_heading_level(number)
                            add_heading(header_text, level)

                        if content:
                            add_paragraph(content, alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

                    elif item_type == "text":
                        if title:
                            add_paragraph(title, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, bold=True,
                                          space_before=Pt(12), space_after=Pt(6))

                        if content:
                            add_paragraph(content, alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

                elif isinstance(item, str):
                    add_paragraph(item, alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

        if json_data.get("conclusion"):
            add_paragraph(json_data["conclusion"], alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                          space_before=Pt(12))

        if json_data.get("appendices"):
            add_paragraph("Appendices:", bold=True, space_before=Pt(12), space_after=Pt(6))
            for item in json_data["appendices"]:
                add_paragraph(f"- {item}", alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, first_line_indent=Cm(0))

        if json_data.get("executor_info"):
            add_paragraph(f"Executor: {json_data['executor_info']}",
                          alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, space_before=Pt(12))

        if json_data.get("distribution_list"):
            add_paragraph("Copy sended:", bold=True, space_before=Pt(12), space_after=Pt(6))
            for target in json_data["distribution_list"]:
                add_paragraph(f"- {target}", alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, first_line_indent=Cm(0))

        if json_data.get("signatures"):
            doc.add_paragraph()

            signatures = json_data["signatures"]
            sender_sig = signatures.get("sender", {})
            recipient_sig = signatures.get("recipient", {})

            sender_text = ""
            if sender_sig:
                label = sender_sig.get("label", "")
                name = sender_sig.get("name", "")
                position = sender_sig.get("position", "")
                sender_text = f"{label}: {position + ' ' if position else ''}{name}"

            recipient_text = ""
            if recipient_sig:
                label = recipient_sig.get("label", "")
                name = recipient_sig.get("name", "")
                position = recipient_sig.get("position", "")
                recipient_text = f"{label}: {position + ' ' if position else ''}{name}"

            if sender_text and not recipient_text:
                add_paragraph(sender_text, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
            elif recipient_text and not sender_text:
                add_paragraph(recipient_text, alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT)
            else:
                add_table_row(left_text=sender_text, right_text=recipient_text)

        if json_data.get("stamp_area"):
            doc.add_paragraph()
            add_paragraph("(STAMP)", alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
        file_path = os.path.join(os.getcwd(), filename)
        doc.save(file_path)
        logger.info(f"Document saved as {file_path}")
        return file_path

    except Exception as e:
        logger.error(f"Error in saving document: {str(e)}", exc_info=True)
        return None


async def process_user_message(message: str, chat_id: int, bot: Bot) -> tuple[str, str | None]:
    global conversation_history

    document_keywords = [
        "contract", "application", "pretense", "power of attorney", "create", "prepare document",
        "complaint", "claim", "receipt", "notification", "act", "agreement", "compose", "make", "draft", "loan",
        "document", "generate"
    ]

    analyze_keywords = [
        "analyze", "check", "evaluation", "analysis", "what's wrong", "remarks", "errors", "is it suitable"
    ]

    conversation_history.append({"role": "user", "content": message})

    document_request_found = any(
        msg["role"] == "user" and not any(kw in msg["content"].lower() for kw in analyze_keywords)
        for msg in conversation_history
    )
    analysis_requested = any(kw in message.lower() for kw in analyze_keywords)
    clarifying_questions_asked = any(
        msg["role"] == "assistant" and "Please, answer the questions:" in msg["content"]
        for msg in conversation_history
    )

    last_uploaded_file = next((msg.get("file_path") for msg in reversed(conversation_history) if msg.get("file_path")),
                              None)

    if analysis_requested and last_uploaded_file:
        analysis_result = analyze_uploaded_document(last_uploaded_file, message)
        response = f"📄 Results of analyzis:\n\n{analysis_result}"
        conversation_history.append({"role": "assistant", "content": response})
        return truncate_if_needed(response), None

    if document_request_found and not clarifying_questions_asked:
        document_result = process_document_request(message, conversation_history)

        if document_result.get("status") == "insufficient_information":
            response = (
                "I lack sufficient information to draft/analyze the document.\n\n"
                f"{document_result.get('explanation')}\n\nPlease answer the following questions:\n"
            )
            for i, question in enumerate(document_result.get("clarifying_questions", []), 1):
                response += f"{i}. {question}\n"

            conversation_history.append({"role": "assistant", "content": response})
            return truncate_if_needed(response), None

        elif document_result.get("status") == "success":
            file_path = save_as_docx(
                document_result.get("document_text"),
                f"{document_result.get('document_type')}_{chat_id}.docx"
            )
            response = f"✅Document «{document_result.get('document_type')}» Ready."
            conversation_history.append({"role": "assistant", "content": response})
            return truncate_if_needed(response), file_path

        else:
            response = (
                f"Failed to create the document. {document_result.get('message', 'Please try rephrasing your request.')}\n\n"
                "Please clarify:\n1. Document type\n2. Main terms\n3. Information about the parties\n4. Additional requests"
            )
            conversation_history.append({"role": "assistant", "content": response})
            return truncate_if_needed(response), None

    if document_request_found and clarifying_questions_asked:
        document_result = generate_document_from_conversation(conversation_history)

        if document_result.get("status") == "success":
            file_path = save_as_docx(
                document_result.get("document_text"),
                f"{document_result.get('document_type')}_{chat_id}.docx"
            )
            response = f"✅ The document «{document_result.get('document_type')}» was successfully created."
            conversation_history.append({"role": "assistant", "content": response})
            return truncate_if_needed(response), file_path

        else:
            response = f"Failed to create the document. {document_result.get('message', 'Please provide more details.')}"
            conversation_history.append({"role": "assistant", "content": response})
            return truncate_if_needed(response), None

    api_messages = [{"role": "system", "content": DOCUMENT_GENERATOR_PROMPT_2}]
    api_messages.extend(conversation_history)

    response = create_openai_completion(messages=api_messages)
    if not response:
        response = "Error."

    conversation_history.append({"role": "assistant", "content": response})
    return truncate_if_needed(response), None


def analyze_uploaded_document(file_path: str, user_message: str) -> str:
    try:
        ext = os.path.splitext(file_path)[-1].lower()

        if ext == ".docx":
            doc = Document(file_path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        elif ext == ".pdf":
            reader = PdfReader(file_path)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)

        else:
            return "❗️ Only files in .docx and .pdf formats are supported."

        prompt = f"""
            The user requests to analyze the document:
            «{user_message}»

            Here is the content of the document:
            {text}

            Please analyze it, point out errors, inconsistencies, legal violations, or recommendations for improvement.
            """
        completion = create_openai_completion([{"role": "user", "content": prompt}])
        return completion or "Failed to obtain document analysis."

    except Exception as e:
        return f"Error reading document: {e}"


def truncate_if_needed(text: str) -> str:
    if len(text) > MAX_TELEGRAM_MESSAGE_LENGTH:
        return text[:MAX_TELEGRAM_MESSAGE_LENGTH - 100] + "\n\n"
    return text


def process_document_request(message: str, conversation_history: List[Dict[str, str]]) -> Dict[str, Any]:
    logger.info(f"ENTER process_document_request(message={message[:50]}...)")

    try:
        analysis_result = create_openai_completion(
            messages=[
                {"role": "system",
                 "content": "You are a legal document analyst. Your task is to analyze requests and determine the necessary information for document preparation."},
                {"role": "user", "content": DOCUMENT_ANALYSIS_PROMPT.format(message=message)},
            ],
            response_format={"type": "json_object"}
        )

        if not analysis_result:
            return {
                "status": "error",
                "message": "Failed to analyze the request."
            }

        try:
            analysis_data = json.loads(analysis_result)
        except json.JSONDecodeError:
            logger.error(f"JSON decode error: {analysis_result}")
            return {
                "status": "error",
                "message": "Failed to process server response."
            }

        if not analysis_data.get("has_sufficient_info", False):
            return {
                "status": "insufficient_information",
                "explanation": "Additional information is required for correct document creation.",
                "clarifying_questions": analysis_data.get("clarifying_questions", [])
            }

        document_type = analysis_data.get("document_type", "legal document")

        document_text = create_openai_completion(
            messages=[
                {"role": "system", "content": DOCUMENT_GENERATOR_PROMPT_2},
                {"role": "user",
                 "content": DOCUMENT_GENERATOR_PROMPT.format(document_type=document_type, message=message)}
            ],
            response_format={"type": "json_object"}
        )

        if not document_text:
            return {
                "status": "error",
                "message": "Failed to generate document text."
            }

        try:
            document_json = json.loads(document_text)
        except json.JSONDecodeError:
            logger.error(f"JSON decode error for document: {document_text}")
            return {
                "status": "error",
                "message": "Failed to process document structure."
            }

        validation_result = validate_document(document_text)
        recommendations = generate_recommendations(document_text)

        return {
            "status": "success",
            "document_type": document_type,
            "document_text": document_json,
            "validation_result": validation_result,
            "recommendations": recommendations
        }

    except Exception as e:
        logger.error(f"Error in process_document_request: {str(e)}", exc_info=True)
        return {
            "status": "error",
            "message": f"An error occurred while processing the request: {str(e)}"
        }


def generate_document_from_conversation(conversation_history: List[Dict[str, str]]) -> Dict[str, Any]:
    logger.info("ENTER generate_document_from_conversation")

    try:
        conversation_text = "\n".join([f"{msg['role'].upper()}: {msg['content']}" for msg in conversation_history])

        document_type = create_openai_completion(
            messages=[
                {"role": "system", "content": "You are a legal document analyst."},
                {"role": "user", "content": DOCUMENT_TYPE_DETECTION_PROMPT.format(conversation_text=conversation_text)}
            ]
        )

        if not document_type:
            document_type = "legal document"

        document_text = create_openai_completion(
            messages=[
                {"role": "system", "content": DOCUMENT_GENERATOR_PROMPT_2},
                {"role": "user", "content": DOCUMENT_GENERATION_FROM_DIALOGUE_PROMPT.format(
                    document_type=document_type,
                    conversation_text=conversation_text
                )}
            ],
            response_format={"type": "json_object"}
        )

        if not document_text:
            return {
                "status": "error",
                "message": "Failed to generate document text."
            }

        try:
            document_json = json.loads(document_text)
        except json.JSONDecodeError:
            logger.error(f"JSON decode error for document: {document_text}")
            return {
                "status": "error",
                "message": "Failed to process document structure."
            }

        validation_result = validate_document(document_text)

        recommendations = generate_recommendations(document_text)

        return {
            "status": "success",
            "document_type": document_type,
            "document_text": document_json,
            "validation_result": validation_result,
            "recommendations": recommendations
        }

    except Exception as e:
        logger.error(f"Error in generate_document_from_conversation: {str(e)}", exc_info=True)
        return {
            "status": "error",
            "message": f"An error occurred while generating the document: {str(e)}"
        }


def validate_document(document_text: str) -> str:
    validation_result = create_openai_completion(
        messages=[
            {"role": "system", "content": "You are an experienced lawyer specializing in document review."},
            {"role": "user", "content": VALIDATION_PROMPT.format(document_text=document_text)}
        ]
    )
    return validation_result


def generate_recommendations(document_text: str) -> str:
    recommendations = create_openai_completion(
        messages=[
            {"role": "system", "content": "You are a legal consultant specializing in practical recommendations."},
            {"role": "user", "content": RECOMMENDATIONS_PROMPT.format(document_text=document_text)}
        ]
    )
    return recommendations


def create_openai_completion(messages: List[Dict[str, str]], response_format: Dict[str, str] = None) -> Optional[str]:
    try:
        logger.info(f"Sending request to OpenAI API with {len(messages)} messages")

        params = {
            "model": "o3-mini",
            "messages": messages,
            "max_completion_tokens": 100000,
        }

        if response_format:
            params["response_format"] = response_format

        response = client.chat.completions.create(**params)

        if response and response.choices and len(response.choices) > 0:
            return response.choices[0].message.content
        else:
            logger.error("Invalid response from OpenAI API")
            return None

    except Exception as e:
        logger.error(f"Error calling OpenAI API: {str(e)}", exc_info=True)
        return None
